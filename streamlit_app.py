import random
import streamlit as st
import gen
st.set_page_config(
    page_title="AVE",
    page_icon="🍊"
)
with st.sidebar:
    api_key = st.text_input("你的api key(默认为空即可)", key="file_qa_api_key", type="password")
    "[狠狠赞助我](https://afdian.net/a/voidfun/plan)"

st.title("📝 AVE --- 新时代Ai英语学习助手")
st.subheader("step1 -你要背的单词")
words = st.text_input("空格隔开即可")
option = st.selectbox(
    '或者选择单词书',
    ('','高中'))
if option =="高中":
    file = open("gz.txt","r").readlines()
    for i in range(20):
        words += file[random.randint(0,1139)]
    st.write(words)
    pass
st.subheader("step2 -选择故事风格(可以多选)")
options = st.multiselect(
    '故事风格',
    ['浪漫', '悬疑', '恐怖', '穿越','古代','幽默','黑色','感动','反转','武侠','童话','科幻','幻想','历史','侦探','成长','哥特','赛博朋克','奇谈怪论','寓言','戏剧','歌剧','意识流'])
style:str=""
if len(options) != 0:
    for i in options:
        if len(style) == 0:
            style=i
        else:
            style+=','+i
genable=False
st.subheader("step3 -选择生成模式")
mode=st.selectbox("你准备生成什么内容",('背诵文章','考试阅读','文章续写'))

inputs=""
if mode=="文章续写":
    inputs=st.text_input("输入你要续写的内容")
    modeable = True if inputs else False
else:
    modeable=True
if words and style and modeable:
    genable=True
st.subheader("step3.5 -可选项")
ex = st.checkbox("生成word文档")
if not ex:
    tl = st.checkbox("听力训练(必须勾选生成文档)",disabled=True)
else:
    tl = st.checkbox("听力训练(没实现)", disabled=True)
lm = st.checkbox("长故事")
st.checkbox("故事视频",disabled=True)
st.subheader("step3 -AI赋予你力量")
writer = st.checkbox("模仿著名作家们的风格?")
if writer:
    types_w = st.multiselect("作家风格",['卡夫卡式','狄更斯式','海明威式','老舍式','多斯妥耶夫斯基式','菲茨杰拉德式','阿加莎·克里斯蒂式','科马克·麦卡锡式','苏珊·柯林斯式'])
if st.button("生成故事",disabled=not genable):
    path = f"{hash(words)}.docx"
    with st.spinner("生成文章中 -请稍候"):
        genable = False
        TEXT = gen.gen(words, option, longmethod=lm,mode=mode,input=inputs)
        if ex:
            with st.spinner("整合文档中"):
                from docx import Document
                import shutil
                shutil.copyfile("temple.docx", path)
                doc = Document(path)
                doc.add_heading(words, level=1)
                doc.add_paragraph(TEXT)
                doc.save(path)
        genable = False
    st.balloons()
    st.info("生成文章成功")
    if ex:
        with open(f"{hash(words)}.docx", "rb") as file:
            btn = st.download_button(
                label="下载",
                data=file,
                file_name="1.doc",
                mime="xml/doc"
            )
    else:
        st.write(TEXT)
